from flask import Flask, render_template, request, redirect, session, send_from_directory,send_file
import sqlite3
import os
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import schedule
import time
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment
from openpyxl.styles import Alignment, Border, Side
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from flask import send_file

app = Flask(__name__)
app.secret_key = os.urandom(24)

# For images
@app.route('/images/<path:filename>')
def serve_images(filename):
    return send_from_directory('images', filename)

# Route for login page
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if check_user(username, password):
            session['username'] = username
            return redirect('/landing')
        else:
            return render_template('login.html', message="Username or password is incorrect.")
    return render_template('login.html', message="")

# Route for signup page
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        if password == confirm_password:
            if not check_username(username):
                conn = sqlite3.connect('users.db')
                c = conn.cursor()
                c.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, password))
                conn.commit()
                conn.close()
                session['username'] = username
                return redirect('/landing')
            else:
                return render_template('signup.html', message="Username already exists. Please choose another username.")
        else:
            return render_template('signup.html', message="Password and confirm password do not match.")
    return render_template('signup.html', message="")

# Route for landing page
@app.route('/landing')
def landing():
    if 'username' in session:
        update_days_remaining()
        conn = sqlite3.connect('medicine.db')
        c = conn.cursor()

        # Get sorting parameters
        sort_by = request.args.get('sort_by', 'name')
        sort_order = request.args.get('sort_order', 'asc')

        # Construct SQL query based on sorting parameters
        if sort_by == 'name':
            order_by = 'name ' + sort_order
        elif sort_by == 'expiry_date':
            order_by = 'expiry_date ' + sort_order
        elif sort_by == 'installation_date':
            order_by = 'installation_date ' + sort_order
        elif sort_by == 'days_remaining':
            order_by = 'days_remaining ' + sort_order
        else:
            order_by = 'name ASC'  # Default sorting

        c.execute("SELECT * FROM medicines WHERE days_remaining > 0 ORDER BY " + order_by)  # Exclude expired medicines
        medicines = c.fetchall()
        conn.close()

        # Calculate number of days until expiry for each medicine
        today = datetime.now().date()

        # Initialize variables to store details of expiring medicines
        expiring_medicines = []
        expired_medicines = []

        for med in medicines:
            if med[2] and med[2].strip():    
                expiry_date = datetime.strptime(med[2], '%Y-%m-%d').date()
                med_days_until_expiry = (expiry_date - today).days
                med = med[:4] + (med_days_until_expiry,) + med[4:]

                # Check if medicine's expiry is within specified range
                if med_days_until_expiry in [10, 5, 4, 3, 2, 1]:
                    expiring_medicines.append(med)

                # Check if medicine has already expired
                if med[4] < 0:
                    expired_medicines.append(med)

            # Check if medicine has already expired
            if med[4] < 0:
                expired_medicines.append(med)

        # Prepare message body for expiring medicines
        expiring_message = ""
        for med in expiring_medicines:
            expiring_message += f"- {med[1]} is expiring in {med[4]} days.\n"

        # Prepare message body for expired medicines
        expired_message = ""
        for med in expired_medicines:
            expired_message += f"- {med[1]} has already expired.\n"

        # Concatenate expiring and expired messages
        message_body = ""
        if expiring_message:
            message_body += "Medicines Expiring Soon:\n" + expiring_message + "\n"
        if expired_message:
            message_body += "Expired Medicines:\n" + expired_message + "\n"

        return render_template('landing.html', medicines=medicines, username=session['username'])

    else:
        return redirect('/')

# Route for expired medicines page
@app.route('/expired_medicines')
def expired_medicines():
    if 'username' in session:
        conn = sqlite3.connect('medicine.db')
        c = conn.cursor()
        c.execute("SELECT * FROM medicines WHERE days_remaining < 0")
        expired_medicines = c.fetchall()
        conn.close()
        return render_template('expired_medicines.html', expired_medicines=expired_medicines, username=session['username'])
    else:
        return redirect('/')

# Route to handle update action for expired medicines
@app.route('/update_expired_medicine/<int:medicine_id>', methods=['POST'])
def update_expired_medicine(medicine_id):
    if 'username' in session:
        if request.method == 'POST':
            name = request.form['medicineName']
            expiry_date = request.form['expiryDate']
            installation_date = request.form['installationDate']
            quantity = request.form['medicineQuantity']
            
            # Update the medicine record in the database
            conn = sqlite3.connect('medicine.db')
            c = conn.cursor()
            c.execute("UPDATE medicines SET name=?, expiry_date=?, installation_date=?, quantity=? WHERE id=?",
                      (name, expiry_date, installation_date, quantity, medicine_id))
            conn.commit()
            
            # Check if the expiry date is ahead
            if expiry_date > datetime.today().strftime('%Y-%m-%d'):
                # Move the medicine from expired to landing page
                c.execute("DELETE FROM expired_medicines WHERE id=?", (medicine_id,))
                conn.commit()
                # Add the medicine to the landing page
                c.execute("INSERT INTO medicines (name, expiry_date, installation_date, quantity) VALUES (?, ?, ?, ?)",
                          (name, expiry_date, installation_date, quantity))
                conn.commit()
                conn.close()
                return redirect('/landing')
            else:
                conn.close()
                return redirect('/expired_medicines')
    else:
        return redirect('/')

# Route to handle delete action for expired medicines
@app.route('/delete_expired_medicine/<int:medicine_id>', methods=['POST'])
def delete_expired_medicine(medicine_id):
    if 'username' in session:
        if request.method == 'POST':
            # Delete the medicine record from the database
            conn = sqlite3.connect('medicine.db')
            c = conn.cursor()
            c.execute("DELETE FROM expired_medicines WHERE id=?", (medicine_id,))
            conn.commit()
            conn.close()
            return redirect('/expired_medicines')
    else:
        return redirect('/')


# Route to add medicine info    
@app.route('/add_info', methods=['POST'])
def add_info():
    if 'username' in session:
        name = request.form['medicineName']
        expiry_date = request.form['expiryDate']
        installation_date = request.form['installationDate']
        quantity = request.form['medicineQuantity']
        conn = sqlite3.connect('medicine.db')
        c = conn.cursor()
        c.execute("INSERT INTO medicines (name, expiry_date, installation_date, quantity) VALUES (?, ?, ?, ?)",
                  (name, expiry_date, installation_date, quantity))
        conn.commit()
        conn.close()

        # Update days remaining for all medicines
        update_days_remaining()

        return redirect('/landing')
    else:
        return redirect('/')

# Route to delete medicine info
@app.route('/delete_info/<int:medicine_id>', methods=['POST'])
def delete_info(medicine_id):
    if 'username' in session:
        conn = sqlite3.connect('medicine.db')
        c = conn.cursor()
        c.execute("DELETE FROM medicines WHERE id=?", (medicine_id,))
        conn.commit()
        conn.close()

        return redirect('/landing')
    else:
        return redirect('/')


# Route to update medicine info
@app.route('/update_info/<int:medicine_id>', methods=['POST'])
def update_info(medicine_id):
    if 'username' in session:
        if request.method == 'POST':
            name = request.form['medicineName']
            expiry_date = request.form['expiryDate']
            installation_date = request.form['installationDate']
            quantity = request.form['medicineQuantity']
            conn = sqlite3.connect('medicine.db')
            c = conn.cursor()
            c.execute("UPDATE medicines SET name=?, expiry_date=?, installation_date=?, quantity=? WHERE id=?",
                      (name, expiry_date, installation_date, quantity, medicine_id))
            conn.commit()
            conn.close()

            # Update days remaining for all medicines
            update_days_remaining()

            return redirect('/landing')
    else:
        return redirect('/')
    
@app.route('/export_to_excel')
def export_to_excel():
    # Fetch medicines from the database
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute("SELECT name, expiry_date, installation_date, quantity, days_remaining FROM medicines where days_remaining > 0 ORDER BY name ASC")
    medicines = c.fetchall()
    conn.close()

    # Create a new Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Medicine Inventory"

    ws.merge_cells('A1:F1')
    main_cell = ws['A1']
    main_cell.alignment = Alignment(horizontal='center', vertical='center')
    main_cell.value = "Active Medicines"
    main_cell.font = Font(bold=True, size=14)
    # Write headers to the first row
    ws.append(["Sr. No", "Name", "Installation Date", "Expiry Date", "Quantity", "Expiring within"])

    # Write data rows
    for idx, med in enumerate(medicines, start=1):
        remaining_days = med[4]
        if remaining_days is not None and remaining_days != '':
            remaining_days = int(remaining_days)
            months = remaining_days // 30
            days = remaining_days % 30
            if months > 0:
                remaining_text = f"{months} months, {days} days"
            elif days > 0:
                remaining_text = f"{days} days"
            else:
                remaining_text = "Medicine Expired"
        else:
            remaining_text = "No expiry date"

        ws.append([idx, med[0], med[2], med[1], med[3], remaining_text])

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center')

    for row_num, row in enumerate(ws.iter_rows(), start=1):
        for cell in row:
            border = Border(left=Side(style='thin'), 
                            right=Side(style='thin'), 
                            top=Side(style='thin'), 
                            bottom=Side(style='thin'))
            cell.border = border

    for col in ws.columns:
        max_length = 1
        column = col[1].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2) * 1.2
        ws.column_dimensions[column].width = adjusted_width

    for _ in range(2):  # Adjust the number of empty rows as needed
        ws.append([])
    
    last_row_index = ws.max_row + 2

    # Calculate the row index of the immediate row after the empty rows
    header_row_index = last_row_index + 1

    # Merge cells for the header row
    ws.merge_cells(start_row=header_row_index, start_column=1, end_row=header_row_index, end_column=6)

    # Access the header cell
    header_cell = ws.cell(row=header_row_index, column=1)

    # Write "Expired Medicines" in the merged cell
    header_cell.value = "Expired Medicines"
    header_cell.alignment = Alignment(horizontal='center', vertical='center')
    header_cell.font = Font(bold=True, size=14)
    ws.append(["Sr. No", "Name", "Installation Date", "Expiry Date", "Quantity", "Expiring within"])
    # Append the medicines whose expiry date is less than zero
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute("SELECT name, expiry_date, installation_date, quantity, days_remaining FROM medicines where days_remaining <= 0 ORDER BY name ASC")
    medicines = c.fetchall()
    conn.close()
    for idx, med in enumerate(medicines, start=1):
        ws.append([idx, med[0], med[2], med[1], med[3], "Expired"])

    # Adjust the border and alignment for the new rows
    max_row = ws.max_row
    for row_num in range(max_row - len(medicines), max_row + 1):
        for cell in ws[row_num]:
            cell.alignment = Alignment(horizontal='center', vertical='center')
            border = Border(left=Side(style='thin'), 
                            right=Side(style='thin'), 
                            top=Side(style='thin'), 
                            bottom=Side(style='thin'))
            cell.border = border

    # Save the workbook to a temporary file
    excel_filename = "medicine_inventory.xlsx"
    excel_path = os.path.join(app.root_path, excel_filename)
    wb.save(excel_path)

    # Return a download link to the Excel file
    return send_from_directory(app.root_path, excel_filename, as_attachment=True)

@app.route('/export_to_word')
def export_to_word():
    # Read the preset header file
    with open('Header.docx', 'rb') as header_file:
        header_doc = Document(header_file)
    
    # Add current date
    current_date = header_doc.add_paragraph()
    current_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_run = current_date.add_run("Date: " + datetime.now().strftime('%d/%m/%Y'))
    date_run.bold = True

    # Fetch active medicines from the database
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute("SELECT name, expiry_date, installation_date, quantity, days_remaining FROM medicines WHERE days_remaining > 0 ORDER BY name ASC")
    active_medicines = c.fetchall()

    # Fetch expired medicines from the database
    c.execute("SELECT name, expiry_date, installation_date, quantity, days_remaining FROM medicines WHERE days_remaining <= 0 ORDER BY name ASC")
    expired_medicines = c.fetchall()

    conn.close()

    # Add active medicines table to the document
    if active_medicines:
        # Add header for active medicines
        act_head = header_doc.add_heading("Active Medicines", level=1)
        act_head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add table for active medicines
        table_active = header_doc.add_table(rows=1, cols=6) 
        table_active.alignment = WD_TABLE_ALIGNMENT.CENTER       

        # Add headers to the active medicines table
        hdr_cells = table_active.rows[0].cells
        hdr_cells[0].text = 'Sr. No'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Installation Date'
        hdr_cells[3].text = 'Expiry Date'
        hdr_cells[4].text = 'Quantity'
        hdr_cells[5].text = 'Expiring within'

        # Add data rows for active medicines
        for idx, med in enumerate(active_medicines, start=1):
            remaining_days = med[4]
            if remaining_days is not None and remaining_days != '':
                remaining_days = int(remaining_days)
                months = remaining_days // 30
                days = remaining_days % 30
                if months > 0:
                    remaining_text = f"{months} months, {days} days"
                elif days > 0:
                    remaining_text = f"{days} days"
                else:
                    remaining_text = "Medicine Expired"
            else:
                remaining_text = "No expiry date"

            row_cells = table_active.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = med[0]
            row_cells[2].text = med[2]
            row_cells[3].text = med[1]
            row_cells[4].text = str(med[3])
            row_cells[5].text = remaining_text

    # Add expired medicines table to the document
    if expired_medicines:
        # Add header for expired medicines
        exp_head = header_doc.add_heading("Expired Medicines", level=1)
        exp_head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add table for expired medicines
        table_expired = header_doc.add_table(rows=1, cols=6)
        table_expired.alignment = WD_TABLE_ALIGNMENT.CENTER       

        # Add headers to the expired medicines table
        hdr_cells = table_expired.rows[0].cells
        hdr_cells[0].text = 'Sr. No'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Installation Date'
        hdr_cells[3].text = 'Expiry Date'
        hdr_cells[4].text = 'Quantity'
        hdr_cells[5].text = 'Status'

        # Add data rows for expired medicines
        for idx, med in enumerate(expired_medicines, start=1):
            row_cells = table_expired.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = med[0]
            row_cells[2].text = med[2]
            row_cells[3].text = med[1]
            row_cells[4].text = str(med[3])
            row_cells[5].text = "Expired"

    # Save the Word document to a temporary file
    word_filename = "medicine_inventory.docx"
    word_path = os.path.join(app.root_path, word_filename)
    header_doc.save(word_path)

    # Return a download link to the Word file
    return send_file(word_path, as_attachment=True)

# Route for logout
@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect('/')    

# Function to create database table
def create_table():
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS medicines
                 (id INTEGER PRIMARY KEY, name TEXT, expiry_date TEXT, installation_date TEXT, quantity INTEGER)''')
    conn.commit()
    conn.close()

# Fetch all medicines from the database
def fetch_medicines():
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute("SELECT * FROM medicines")
    medicines = c.fetchall()
    conn.close()

# Update Days to Expiry
def update_days_remaining():
    conn = sqlite3.connect('medicine.db')
    c = conn.cursor()
    c.execute("SELECT id, expiry_date FROM medicines")
    medicines = c.fetchall()

    # Get current date
    today = datetime.now().date()

    # Update days remaining for each medicine
    for med_id, expiry_date in medicines:
        if expiry_date and expiry_date.strip():
            expiry_date = datetime.strptime(expiry_date, "%Y-%m-%d").date()
            days_remaining = (expiry_date - today).days
        else:
            days_remaining = ''
        c.execute("UPDATE medicines SET days_remaining = ? WHERE id = ?", (days_remaining, med_id))

    # Commit changes and close connection
    conn.commit()
    conn.close()

# Function to send email
def send_mail(subject, message, to_email):
    from_email = 'itfakms@gmail.com'
    password = 'mlnk pftp ejui jznw'

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ', '.join(to_email)
    msg['Subject'] = subject
    msg.attach(MIMEText(message, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
    except Exception as e:
        print("Error sending email:", str(e))

# Function to check and send email at 11 PM
def check_and_send_email():
    current_time = datetime.now().time()
    if current_time.hour == 15   and current_time.minute == 50:
        print("Checking medicines and sending email...")
        # Update days remaining for all medicines
        update_days_remaining()

        # Fetch all medicines
        conn = sqlite3.connect('medicine.db')
        c = conn.cursor()
        c.execute("SELECT * FROM medicines")
        medicines = c.fetchall()
        conn.close()

        # Calculate number of days until expiry for each medicine
        today = datetime.now().date()

        # Initialize variables to store details of expiring medicines
        expiring_medicines = []
        expired_medicines = []

        for med in medicines:
            expiry_date = med[2]
            if expiry_date and expiry_date.strip():
                expiry_date = datetime.strptime(med[2], '%Y-%m-%d').date()
                med_days_until_expiry = (expiry_date - today).days
                med = med[:4] + (med_days_until_expiry,) + med[4:]

                # Check if medicine's expiry is within specified range
                if med_days_until_expiry in [10, 5, 4, 3, 2, 1]:
                    expiring_medicines.append(med)

                # Check if medicine has already expired
                if med[4] < 0:
                    expired_medicines.append(med)

        # Prepare message body for expiring medicines
        expiring_message = ""
        for med in expiring_medicines:
            expiring_message += f"- {med[1]} is expiring in {med[4]} days.\n"

        # Prepare message body for expired medicines
        expired_message = ""
        for med in expired_medicines:
            expired_message += f"- {med[1]} has already expired.\n"

        # Concatenate expiring and expired messages
        message_body = ""
        if expiring_message:
            message_body += "Medicines Expiring Soon:\n" + expiring_message + "\n"
        if expired_message:
            message_body += "Expired Medicines:\n" + expired_message + "\n"

        # Send email with concatenated message body
        if message_body:
            send_mail("Medicine Expiry Alert", message_body, ['ashishjoshi2021.it@mmcoe.edu.in', 'atharvaphadke2021.it@mmcoe.edu.in', 'soahammohaadkar2021.it@mmcoe.edu.in'])
 
        print("Email sent.")

# Schedule email check every minute
schedule.every().day.at("15:50").do(check_and_send_email)

# Run the scheduler in a separate thread
def scheduler_thread():
    while True:
        schedule.run_pending()
        time.sleep(1)

# Function to create database table
def create_table():
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY, username TEXT, password TEXT)''')
    conn.commit()
    conn.close()

# Function to check if username and password match
def check_user(username, password):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE username=?", (username,))
    user = c.fetchone()
    conn.close()
    if user and user[2] == password:
        return True
    return False

# Function to check if username exists
def check_username(username):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE username=?", (username,))
    user = c.fetchone()
    conn.close()
    return user

if __name__ == '__main__':
    create_table()
    # Start scheduler thread
    import threading
    threading.Thread(target=scheduler_thread, daemon=True).start()
    app.run(debug=True)